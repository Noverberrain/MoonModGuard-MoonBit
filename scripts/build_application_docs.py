from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    KeepTogether,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "competition"
PDF_PATH = OUT_DIR / "MoonModGuard项目申报书.pdf"
DOCX_PATH = OUT_DIR / "MoonModGuard项目申报书.docx"

TITLE = "MoonModGuard 开源技术说明书"
SUBTITLE = "MoonBit 项目清单与供应链策略审计器"
GITHUB_URL = "https://github.com/Noverberrain/MoonModGuard-MoonBit-"
GITLINK_URL = "https://gitlink.org.cn/Wyc060514/moonmodguard"
CORE_AUTHOR = "wyc060514"

HIGHLIGHTS = [
    ("定位", "面向 MoonBit 工程清单的发布前自检与策略审计工具。"),
    ("对象", "解析 moon.mod、moon.pkg，提取元数据、依赖和导入别名。"),
    ("价值", "补齐清单完整性、许可证策略和依赖透明度检查能力。"),
    ("交付", "提供核心库、CLI 示例、测试、CI、README 与申报材料。"),
]

BODY_SECTIONS = [
    (
        "评审摘要",
        "MoonModGuard 面向 MoonBit 项目的工程清单质量与供应链策略检查，提供轻量、可测试、可嵌入的审计内核。它解析 moon.mod 与 moon.pkg 文本，抽取项目元数据、包级依赖、导入别名和发布信息，并将这些信息统一转化为可复查的策略诊断报告。首版优先补齐发布前自检能力，而不是构建平台型服务。",
    ),
    (
        "选题价值",
        "MoonBit 生态会持续积累包、模板和教学工程。随着项目数量增加，单个仓库是否具备可发布、可维护、可审计的基础条件，会直接影响生态复用效率。MoonModGuard 将人工阅读清单文件的经验沉淀为基础软件能力，适用于包发布、竞赛验收、课程规范检查和后续 CI 策略审计。",
    ),
    (
        "技术方案",
        "系统采用清单解析层、项目建模层、策略评估层和报告生成层。解析层读取 moon.mod 与 moon.pkg 字符串；建模层合并模块级与包级信息；策略层检查必填元数据、许可证 allowlist、可信依赖前缀和重复依赖；报告层输出稳定 Markdown，供 CLI、CI 或文档系统复用。",
    ),
    (
        "实现边界",
        "首版采用显式输入字符串与内存快照的方式，不做真实文件系统扫描、在线包仓库查询或完整 MoonBit 语法解析。这个边界降低实现风险，使核心逻辑更容易测试，也便于后续扩展真实工作区扫描器。",
    ),
    (
        "验证证据",
        "当前验证覆盖 moon.mod 标量字段解析、keywords 数组解析、模块级 import 解析、包级 import 解析、缺失元数据诊断、非 allowlist 许可证诊断、未知依赖前缀诊断、重复依赖诊断和 Markdown 报告稳定输出。验证命令包括 moon info、moon fmt --check、moon test 和 moon run cmd/main。",
    ),
    (
        "开源计划",
        "项目以 GitHub 仓库作为官方开源发布地址，满足赛事阶段一对有效 GitHub 仓库链接的要求；GitLink 仓库作为赛事平台镜像，用于同步提交和平台验收。后续迭代将优先保持核心 API 稳定、测试可复现、报告输出可比对。",
    ),
]

MODULES = [
    ("parse_mod", "解析模块清单中的名称、版本、许可证、README、仓库、关键词和依赖。"),
    ("parse_pkg", "解析包级 import 依赖与别名。"),
    ("project_from", "组合模块与包信息，构建统一项目模型。"),
    ("evaluate_policy", "依据策略生成元数据、许可证和依赖诊断。"),
    ("render_markdown", "输出稳定审计文本，便于提交、比对和 CI 集成。"),
]


def register_font() -> str:
    for path in [
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]:
        if path.exists():
            pdfmetrics.registerFont(TTFont("MoonModGuardCN", str(path)))
            return "MoonModGuardCN"
    return "Helvetica"


def p(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(text.replace("\n", "<br/>"), style)


def build_pdf() -> None:
    font = register_font()
    styles = getSampleStyleSheet()
    title = ParagraphStyle(
        "TitleCN",
        parent=styles["Title"],
        fontName=font,
        fontSize=24,
        leading=31,
        alignment=TA_LEFT,
        textColor=colors.HexColor("#16324f"),
        spaceAfter=8,
    )
    subtitle = ParagraphStyle(
        "SubtitleCN",
        parent=styles["BodyText"],
        fontName=font,
        fontSize=12,
        leading=18,
        textColor=colors.HexColor("#4a6178"),
        spaceAfter=14,
    )
    label = ParagraphStyle(
        "LabelCN",
        parent=styles["BodyText"],
        fontName=font,
        fontSize=8.5,
        leading=11,
        textColor=colors.HexColor("#5f6f80"),
        alignment=TA_CENTER,
    )
    card_text = ParagraphStyle(
        "CardTextCN",
        parent=styles["BodyText"],
        fontName=font,
        fontSize=9.2,
        leading=14,
        textColor=colors.HexColor("#1f2d3a"),
        alignment=TA_CENTER,
    )
    heading = ParagraphStyle(
        "HeadingCN",
        parent=styles["Heading2"],
        fontName=font,
        fontSize=13.5,
        leading=19,
        textColor=colors.HexColor("#16324f"),
        spaceBefore=9,
        spaceAfter=5,
    )
    body = ParagraphStyle(
        "BodyCN",
        parent=styles["BodyText"],
        fontName=font,
        fontSize=10.3,
        leading=17,
        firstLineIndent=20,
        textColor=colors.HexColor("#202a33"),
        spaceAfter=6,
    )
    small = ParagraphStyle(
        "SmallCN",
        parent=styles["BodyText"],
        fontName=font,
        fontSize=9.2,
        leading=13,
        textColor=colors.HexColor("#263544"),
    )
    mono = ParagraphStyle(
        "MonoCN",
        parent=small,
        fontName=font,
        textColor=colors.HexColor("#16324f"),
    )

    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        rightMargin=1.75 * cm,
        leftMargin=1.75 * cm,
        topMargin=1.45 * cm,
        bottomMargin=1.45 * cm,
        title=TITLE,
    )

    story = [
        p("MOONBIT OPEN SOURCE BRIEF", label),
        p(TITLE, title),
        p(SUBTITLE, subtitle),
    ]

    meta_rows = [
        ["GitHub 主仓库", GITHUB_URL],
        ["GitLink 镜像仓库", GITLINK_URL],
        ["核心作者", CORE_AUTHOR],
        ["开源许可证", "Apache-2.0"],
        ["参赛方向", "MoonBit 国产基础软件开源生态项目"],
    ]
    meta_table = Table(
        [[p(k, small), p(v, mono)] for k, v in meta_rows],
        colWidths=[3.1 * cm, 12.2 * cm],
    )
    meta_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f5f8fb")),
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#e7eef5")),
                ("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor("#9eb2c5")),
                ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#c8d4df")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.extend([meta_table, Spacer(1, 0.28 * cm)])

    highlight_cells = [
        [p(name, label), p(text, card_text)] for name, text in HIGHLIGHTS
    ]
    highlight_table = Table(
        [
            [highlight_cells[0], highlight_cells[1]],
            [highlight_cells[2], highlight_cells[3]],
        ],
        colWidths=[7.45 * cm, 7.45 * cm],
    )
    highlight_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#ffffff")),
                ("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor("#b7c6d6")),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#d4dee8")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
            ]
        )
    )
    story.extend([highlight_table, Spacer(1, 0.22 * cm)])

    for section_title, text in BODY_SECTIONS:
        story.append(KeepTogether([p(section_title, heading), p(text, body)]))

    story.append(p("核心 API 交付", heading))
    module_table = Table(
        [[p(name, mono), p(desc, small)] for name, desc in MODULES],
        colWidths=[3.4 * cm, 11.5 * cm],
    )
    module_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#eef4f8")),
                ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#b9c7d5")),
                ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#d5dee7")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(module_table)
    doc.build(story)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_rule(paragraph, color: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def add_run(paragraph, text: str, size: int, color: str, bold: bool = False):
    run = paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    return run


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(kicker, "MOONBIT OPEN SOURCE BRIEF", 8, "6A7886", True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(title, TITLE, 24, "16324F", True)
    set_paragraph_rule(title, "9EB2C5")

    sub = doc.add_paragraph()
    add_run(sub, SUBTITLE, 12, "4A6178")
    sub.paragraph_format.space_after = Pt(10)

    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.style = "Table Grid"
    meta_rows = [
        ("GitHub 主仓库", GITHUB_URL),
        ("GitLink 镜像仓库", GITLINK_URL),
        ("核心作者", CORE_AUTHOR),
        ("开源许可证", "Apache-2.0"),
        ("参赛方向", "MoonBit 国产基础软件开源生态项目"),
    ]
    for row, (key, value) in zip(meta_table.rows, meta_rows):
        set_cell_shading(row.cells[0], "E7EEF5")
        set_cell_shading(row.cells[1], "F7FAFC")
        row.cells[0].text = key
        row.cells[1].text = value

    doc.add_paragraph()
    highlight_table = doc.add_table(rows=2, cols=2)
    highlight_table.style = "Table Grid"
    for cell, (name, text) in zip([c for r in highlight_table.rows for c in r.cells], HIGHLIGHTS):
        set_cell_shading(cell, "FFFFFF")
        cell.text = ""
        para = cell.paragraphs[0]
        add_run(para, f"{name}\n", 9, "5F6F80", True)
        add_run(para, text, 9, "1F2D3A")

    for section_title, text in BODY_SECTIONS:
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(9)
        h.paragraph_format.space_after = Pt(3)
        add_run(h, section_title, 14, "16324F", True)
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = Pt(21)
        para.paragraph_format.line_spacing = 1.35
        add_run(para, text, 10, "202A33")

    h = doc.add_paragraph()
    add_run(h, "核心 API 交付", 14, "16324F", True)
    module_table = doc.add_table(rows=len(MODULES), cols=2)
    module_table.style = "Table Grid"
    for row, (name, desc) in zip(module_table.rows, MODULES):
        set_cell_shading(row.cells[0], "EEF4F8")
        row.cells[0].text = name
        row.cells[1].text = desc

    doc.save(DOCX_PATH)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_pdf()
    build_docx()
    print(PDF_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
